# Libraries.
# Pandas for reading excel and transformations.
import pandas as pd
# Openpyxl for the reading engine and extracting the ticket embeded URLs
import openpyxl
# Docx for creating the word files.
import docx
# Necessary to import in order to define the image dimensions
from docx.shared import Inches
# Necessary to import for the hyperlink formatting.
from docx.enum.dml import MSO_THEME_COLOR_INDEX
# OS for checking, deleting and creating files and folders.
import os
# Win32 for the interaction with Outlook. Creating and sending emails.
import win32com.client as win32
# Plotly for the Treemap plot.
import plotly.express as px

# Check for temp folder and create.
if not os.path.isdir("temp"):
    os.mkdir("temp")
    
# Path to the Audit extraction excel file.
AuExpath = ""

# Path to the dataset owner extraction csv file.
dfOwpath = ""

# Path to the Corrections file. Including the Dataset Owners email. Manually maintained.
Corpath = ""

# Reading of the extraction file.
dfAuEx_all = pd.read_excel(AuExpath, header=0, thousands='.', engine='openpyxl')

# Reading of the dataset owner extraction file.
dfOw_all = pd.read_csv(dfOwpath, skiprows=0, quotechar='"', header=1, delimiter=',', engine='python', usecols=[1,2], names=['Dataset','Data_Owner'])

# Hyperlink extraction from the ticket number column from the Audit extraction excel file.
# Taken with modifications from: "https://stackoverflow.com/a/68031474"
wb = openpyxl.load_workbook(AuExpath)
ws = wb['Data Quality - Extr']
Tlink = []
for col in ws['E']:
    try:
        Tlink.append(col.hyperlink.target)
    except:
        Tlink.append(col.value)
        
# Insertion of the hyperlink column into the Audit extraction dataframe.
dfAuEx_all.insert(loc=5,column='Field_3_Hyperlink', value=Tlink[1:])

# Reading the Corrections file, worksheet Config. Can be used to input the path variables and Display or Send. NOT USED.
#Conf_all = pd.read_excel(Corpath, sheet_name='Config' , header=0, engine='openpyxl')
# Reading the Corrections file, worksheet Data_Owners.
Addr_all = pd.read_excel(Corpath, sheet_name='Data_Owners' , header=0, engine='openpyxl')
# Reading the Corrections file, worksheet Corrections.
Cor_all = pd.read_excel(Corpath, sheet_name='Corrections' , header=0, engine='openpyxl')

# Renaming of the columns substituting spaces for undescores
names = dfAuEx_all.columns.values.tolist()
names = [s.replace(' ', '_') for s in names]
dfAuEx_all.columns = names

# Renaming of the columns substituting spaces for undescores
names = dfOw_all.columns.values.tolist()
names = [s.replace(' ', '_') for s in names]
dfOw_all.columns = names

# Setting the index for the update
Cor_all = Cor_all.set_index('Ticket_Number', drop=False)
dfAuEx_all = dfAuEx_all.set_index('Ticket_Number', drop=False)

# Update to change the dataset corrections and reset index.
dfAuEx_all.update(Cor_all)
dfAuEx_all.reset_index(inplace=True, drop=True)

# To add another line to the dfOw_all. Can be reused to manually introduce Datasets and Data_Owners.
vec1 = pd.Series({'Field_1': 'Value_A', 'Field_2': 'Value_B'})
dfOw_all = pd.concat([dfOw_all, vec1.to_frame().T], ignore_index=True)
# Change the dataset owner from "" to "" to aggregate the report.
dfOw_all['Field_2'].replace('Old_Value','New_Value',inplace=True)

# Creation of new column extracting the Customer directorate from the complete one.
dfAuEx_all['Field_4_Extract'] = dfAuEx_all['Field_4'].str.extract('\/(.*?)\/', expand=True)
# Extract from the Field_2 column. Just the starting until the /
dfOw_all['Field_2_Extract'] = dfOw_all['Field_2'].str.extract('(.*?)\/', expand=True)

# Filter the open ones.
open_tickets = ['In Progress', 'Assigned', 'Pending', 'Waiting Approval', 'Planning']
#closed_tickets = ['Closed', 'Completed']

dfAuEx_all = dfAuEx_all[
    (dfAuEx_all.DQ_Relation == 'Yes') &
    (dfAuEx_all.Status.isin(open_tickets))
    & ~(dfAuEx_all.Organization == 'Org_1')
    # & (dfAuEx_all.Status.isin(closed_tickets))
    # & (dfAuEx_all.Submit_Date > '2023-03-01')
].sort_values(['Ticket_Number'])

# Merge extraction with Owners
dfOw_all = dfOw_all.merge(Addr_all, left_on='Data_Owner', right_on='Data_Owner', how='left')

# Merge of dfAuEx_all and dfOw_all by Dataset
dfAuEx_all['Dataset_lower'] = dfAuEx_all['Dataset'].str.lower()
dfOw_all['Dataset_lower'] = dfOw_all['Dataset'].str.lower()
df_all = dfAuEx_all.merge(dfOw_all, left_on='Dataset_lower', right_on='Dataset_lower', suffixes=('','_OW'), how='left')

# Fill empty Dataset and Data_Owner_Directorate with 'Missing'
df_all["Dataset"] = df_all["Dataset"].fillna('Missing')
df_all["Data_Owner_Directorate"] = df_all["Data_Owner_Directorate"].fillna('Missing')
df_all["Data_Owner"] = df_all["Data_Owner"].fillna('Missing')

# Definition of functions for report

# Function that builds the list and details of the subset tickets.
def wordBuild(word_doc, df_subset, dataS):
    # Adds the Ticket number as hyperlink.
    add_hyperlink(word_doc.add_heading(),"Ticket number: " + df_subset.iloc[dataS].Ticket_Number, df_subset.iloc[dataS].Ticket_Hlink)
    # Ticket title (Summary).
    word_doc.add_paragraph().add_run(df_subset.iloc[dataS].Summary, style = "Intense Emphasis")
    # Dataset, Nbr_Days, Status, Description. Other columns can be included modifying the string "" and the query (df_subset...).
    word_doc.add_paragraph("Dataset: ").add_run(df_subset.iloc[dataS].Dataset).bold = True
    word_doc.add_paragraph("Nbr_Days: ").add_run(str(df_subset.iloc[dataS].Nbr_Days)).bold = True
    word_doc.add_paragraph("Priority: ").add_run(df_subset.iloc[dataS].Priority).bold = True
    word_doc.add_paragraph("Status: ").add_run(df_subset.iloc[dataS].Status).bold = True
    word_doc.add_paragraph().add_run("Description:").bold = True
    word_doc.add_paragraph(df_subset.iloc[dataS].Description)
    if not pd.isnull(df_subset.iloc[dataS].Root_cause):
        word_doc.add_paragraph().add_run("Root_cause").bold = True
        word_doc.add_paragraph(str(df_subset.iloc[dataS].Root_cause))
    if not pd.isnull(df_subset.iloc[dataS].Impact_analysis):
        word_doc.add_paragraph().add_run("Impact_analysis:").bold = True
        word_doc.add_paragraph(str(df_subset.iloc[dataS].Impact_analysis))
    if not pd.isnull(df_subset.iloc[dataS].Next_steps):
        word_doc.add_paragraph().add_run("Next_Steps").bold = True
        word_doc.add_paragraph(str(df_subset.iloc[dataS].Next_steps))
    pass
    
# Subsets the main dataframe df_all ordered by number of days and creates the fig object that contains the treemap plot.
def NbrPlot(Dir):
    # Subset the table by the Data Owner, order descending by the Nbr. days
    df_subset = df_all[(df_all.Dataset == Dir)].sort_values(["Nbr_Days"], ascending=False)
    # Replaces in the tickets that have 0 Nbr. of days by 1, since unable to represent a 0 area in the plot.
    df_subset["Nbr_Days"].replace(to_replace = 0, value = 1, inplace=True)
    # Treemap object to build the plot.
    fig = px.treemap(df_subset, path=[px.Constant(df_subset.iloc[0].Dataset + ' DQ Tickets Avrg. Nbr. of days: ' + str("%.0f" % df_subset.Nbr_Days.mean())),
                                  'Dataset','Ticket_Number'], values='Nbr_Days', color='Nbr_Days', color_continuous_scale='RdBu_r', 
                     color_continuous_midpoint=df_subset.Nbr_Days.mean())
    fig.update_traces(root_color="white", textinfo="label+text+value", selector=dict(type='treemap'),
                 texttemplate='%{label}<br>Nbr. of days: %{value}', marker=dict(cornerradius=5))
    fig.update_layout(margin = dict(t=0, l=0, r=0, b=0),coloraxis_showscale=False)
    fig.write_image("temp/" + Dir + ".png", scale=3)

# Builds the word_doc used to create the word file and saves it with the database name.
def single_doc(Dir):
    # Subset the table by the Data Owner, order descending by the Nbr. days
    df_subset = df_all[(df_all.Dataset == Dir)].sort_values(["Nbr_Days"], ascending=False)
    # Create docx object
    word_doc = docx.Document()
    # Add heading
    word_doc.add_heading("Dataset: " + str(df_subset.iloc[0].Dataset), 0)
    
    # Add the treemap plot as a picture.
    word_doc.add_heading("Distribution by Dataset and Nbr. of Days", 1)
    word_doc.add_picture("temp/" + Dir + ".png", width=Inches(7.0))
    
    # Add number of open tickets by subset
    word_doc.add_heading("Nbr of open tickets: " + str(df_subset.shape[0]), 1)
    
    # Column subset for the summary. If necessary add new columns.
    # df_summaryL includes the link.
    df_summaryL = df_subset.filter(['Ticket_Number','Ticket_Hlink','Summary','Dataset','Priority','Status', 'Request_Assignee','Nbr_Days'], axis=1)
    df_summary = df_summaryL.drop('Ticket_Hlink', axis=1)
    # Add the empty summary table
    summary_table = word_doc.add_table(df_summary.shape[0]+1, df_summary.shape[1], style="Medium Grid 2 Accent 1")
    # Fill the summary table with the summary dataframe
    for j in range(df_summary.shape[-1]):
        summary_table.cell(0,j).text = df_summary.columns[j]
    for i in range(df_summary.drop('Ticket_Number', axis=1).shape[0]):
        for j in range(df_summary.drop('Ticket_Number', axis=1).shape[-1]):
            summary_table.cell(i+1,j+1).text = str(df_summary.drop('Ticket_Number', axis=1).values[i,j])
            
    for i in range(df_summary.shape[0]):
        row_cells = summary_table.cell(i+1,0).add_paragraph()
        add_hyperlink(row_cells, df_summaryL.values[i,0], df_summaryL.values[i,1])
   
    # Execute the wordBuild function
    for i in range(len(df_subset)):
        wordBuild(word_doc, df_subset, i)
    
    # Save the docx object as a word file
    word_doc.save("temp/" + df_subset.iloc[0].Dataset + ".docx")
    
def mailProd(Own):
    # Subset of df_all by single Data_Owner.
    df_subset = df_all[(df_all.Data_Owner == Own)].sort_values(["Nbr_Days"], ascending=False)
    mail = win32.Dispatch('outlook.application').CreateItem(0)
    # Email recipiant. Data_Owner email as on Corrections file.
    mail.To = df_subset.iloc[0].Owner_email
    # To introduce CC on the email. If multiple separate by ';' in a new column in corrections by the Data_Owner_email named "CC_email"
    #mail.CC = df_subset.iloc[0].CC_email
    # Email subject
    mail.Subject = ''.join(['Dataset owner report ', str(df_subset.iloc[0].Data_Owner)])
    # Email text body in html.
    mail.HTMLBody = ''.join([
        '<p>Dear ',str(df_subset.iloc[0].Owner_name),',<br></p>'
        '<p>As announced (<a href="https://">link</a>), ',
        '<p>Please have a look at the DQ issues attached and come back to us if you have ',
        'any doubt/questions on their status or if you think that any of the issues mentioned '])
    
    # For loop attaching all Dataset word files to their corresponding owner email
    for dtset in df_subset.Dataset.unique():
        mail.Attachments.Add(os.path.abspath("temp/" + dtset + ".docx"))
    
    # Display opens the email window (False prevents the script to pause while open unsent email windows).
    mail.Display(False)
    # Send automatically sends the email without confirmation.
    #mail.Send
    
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink
    
# For loop on dataset to operate the plot function (NbrPlot) and Word files creation (single_doc).
for i in df_all.Dataset.unique():
    NbrPlot(i)
    single_doc(i)
    
# For loop on Data_Owner to operate the email creation (mailProd).
for i in df_all.Data_Owner.unique():
    mailProd(i)
    
# Creation of an excel file 
with pd.ExcelWriter('Ticket_missing_ownership.xlsx') as writer:
    df_all[(df_all.Dataset == "Missing") | (df_all.Dataset == "Other")].sort_values(['Ticket_Number']).to_excel(writer, index=False, sheet_name='Missing-other_dataset')
    df_all[(df_all.Data_Owner_Directorate == "Missing")].sort_values(['Ticket_Number']).to_excel(writer, index=False, sheet_name='Missing_owner')
    pd.DataFrame(pd.Series(df_all.loc[df_all.Owner_email.isna() & df_all.Data_Owner.notna() & ~(df_all.Data_Owner=='Missing')]['Data_Owner'].unique(), name='Missing_owner_email')).to_excel(writer, index=False, sheet_name='missing_owner_email')
    df_all.sort_values(['Data_Owner','Dataset']).groupby(['Dataset','Data_Owner'], as_index=False).size().to_excel(writer, index=False, sheet_name='Datasets-Owner-Nbr.Tickets')
    df_all.filter(['Ticket_Number','Dataset','Data_Owner']).to_excel(writer, index=False, sheet_name='Tickets-Datasets-Owners')
    
# Deletes the contents of the temp folder.
[os.remove(os.path.join("temp", file)) for file in os.listdir("temp")]
# Deletes the temp folder.
os.rmdir("temp")
